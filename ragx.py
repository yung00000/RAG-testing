import os
import glob
import logging
import PyPDF2
import docx
import pandas as pd
import chromadb
import psutil
from flask import Flask, render_template, request, jsonify
from llama_index.core import VectorStoreIndex, Settings, PromptTemplate, Document
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama
from langchain.text_splitter import RecursiveCharacterTextSplitter

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[logging.FileHandler('rag_system.log'), logging.StreamHandler()]
)
logger = logging.getLogger(__name__)

# Paths
KNOWLEDGE_BASE_PATH = r"C:\rag_system\knowledge_base"
VECTOR_DB_PATH = r"C:\rag_system\vector_db"
os.makedirs(KNOWLEDGE_BASE_PATH, exist_ok=True)
os.makedirs(VECTOR_DB_PATH, exist_ok=True)

# Flask app
app = Flask(__name__)
index = None

def check_system_resources():
    """Check available system memory."""
    mem = psutil.virtual_memory()
    return {'available_ram': mem.available / 1024**3}

def select_model():
    """Select a lightweight LLM suitable for 20GB RAM."""
    models = [("qwen:4b", 7)]  # Qwen 4B, optimized for low memory
    resources = check_system_resources()
    
    for model, size in models:
        if resources['available_ram'] > size * 1.5:
            logger.info(f"Selected model: {model}")
            return model
    logger.error("No suitable model found. Ensure Qwen 4B is installed.")
    exit(1)

def load_documents(folder_path, max_files=10):
    """Load and extract text from PDF, DOCX, and XLSX files."""
    file_patterns = ["*.pdf", "*.docx", "*.xlsx"]
    documents, metadatas = [], []
    
    for pattern in file_patterns:
        files = glob.glob(os.path.join(folder_path, pattern))[:max_files]
        for file in files:
            try:
                text = ""
                if file.endswith(".pdf"):
                    with open(file, "rb") as f:
                        reader = PyPDF2.PdfReader(f)
                        text = "".join(page.extract_text() or "" for page in reader.pages[:50])
                elif file.endswith(".docx"):
                    doc = docx.Document(file)
                    text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
                elif file.endswith(".xlsx"):
                    df = pd.read_excel(file)
                    text = df.to_string(index=False)
                
                if text.strip():
                    documents.append(Document(text=text))
                    metadatas.append({"source": os.path.basename(file)})
                    logger.info(f"Loaded: {os.path.basename(file)}")
            except Exception as e:
                logger.error(f"Error processing {file}: {str(e)}")
    
    return documents, metadatas

def split_text(documents, metadatas):
    """Split documents into smaller chunks for indexing."""
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=300, chunk_overlap=50)
    chunks, chunk_metas = [], []
    
    for doc, meta in zip(documents, metadatas):
        split_chunks = text_splitter.split_text(doc.text)
        for i, chunk in enumerate(split_chunks):
            chunks.append(Document(text=chunk))
            chunk_metas.append({**meta, "chunk_id": i})
    
    logger.info(f"Text splitting complete: {len(chunks)} chunks")
    return chunks, chunk_metas

def build_index(documents, metadatas):
    """Build vector index from document chunks."""
    # Initialize ChromaDB
    chroma_client = chromadb.PersistentClient(path=VECTOR_DB_PATH)
    try:
        chroma_client.delete_collection("knowledge_base")
    except:
        pass
    chroma_collection = chroma_client.create_collection("knowledge_base")
    vector_store = ChromaVectorStore(chroma_collection=chroma_collection)
    
    # Setup models
    Settings.llm = Ollama(model=select_model(), request_timeout=300)
    Settings.embed_model = HuggingFaceEmbedding(
        model_name="BAAI/bge-small-zh-v1.5",  # Lightweight embedding model
        embed_batch_size=16
    )
    
    # Create index
    return VectorStoreIndex(
        documents,
        vector_store=vector_store,
        embed_model=Settings.embed_model
    )

def initialize_system():
    """Initialize the RAG system."""
    global index
    logger.info("Initializing RAG system...")
    
    docs, metas = load_documents(KNOWLEDGE_BASE_PATH)
    if not docs:
        logger.error("No documents loaded")
        return False
    
    chunks, chunk_metas = split_text(docs, metas)
    index = build_index(chunks, chunk_metas)
    logger.info("System initialization complete")
    return True

@app.route('/')
def home():
    """Render the main page."""
    resources = check_system_resources()
    return render_template(
        'index_zh_tw.html',
        model_name=select_model(),
        ram_available=f"{resources['available_ram']:.1f}GB"
    )

@app.route('/query', methods=['POST'])
def handle_query():
    """Handle user queries."""
    if not index:
        return jsonify({"error": "System not initialized"}), 503
    
    data = request.get_json()
    query = data.get('query', '').strip()
    if not query:
        return jsonify({"error": "Empty query"}), 400
    
    try:
        query_engine = index.as_query_engine(
            similarity_top_k=2,
            text_qa_template=PromptTemplate(
                "你是一個專門回答繁體中文問題的AI助手，請根據以下上下文用臺灣正體中文回答。\n"
                "上下文:\n{context_str}\n問題: {query_str}\n回答:"
            )
        )
        response = query_engine.query(query)
        return jsonify({"response": str(response)})
    except Exception as e:
        logger.error(f"Query error: {str(e)}")
        return jsonify({"error": "Query processing failed"}), 500

def create_template():
    """Create HTML template for the frontend."""
    template_dir = os.path.join(os.path.dirname(__file__), 'templates')
    os.makedirs(template_dir, exist_ok=True)
    
    with open(os.path.join(template_dir, 'index_zh_tw.html'), 'w', encoding='utf-8') as f:
        f.write('''<!DOCTYPE html>
<html lang="zh-TW">
<head>
    <meta charset="UTF-8">
    <title>繁體中文知識庫</title>
    <style>
        body { font-family: 'Microsoft JhengHei', Arial, sans-serif; padding: 20px; background: #f0f4f8; }
        .container { max-width: 800px; margin: 0 auto; background: white; padding: 20px; border-radius: 8px; }
        textarea { width: 100%; height: 100px; padding: 10px; margin-bottom: 10px; }
        button { background: #2b6cb0; color: white; padding: 10px 20px; border: none; cursor: pointer; }
        button:hover { background: #2c5282; }
        #response { margin-top: 20px; padding: 10px; border: 1px solid #ddd; }
    </style>
</head>
<body>
    <div class="container">
        <h1>繁體中文知識庫</h1>
        <p>模型: {{ model_name }} | 可用記憶體: {{ ram_available }}</p>
        <textarea id="query" placeholder="輸入問題..."></textarea>
        <button onclick="submitQuery()">查詢</button>
        <div id="response">請輸入問題</div>
    </div>
    <script>
        function submitQuery() {
            const query = document.getElementById('query').value.trim();
            if (!query) { alert('請輸入問題'); return; }
            fetch('/query', {
                method: 'POST',
                headers: {'Content-Type': 'application/json'},
                body: JSON.stringify({ query })
            })
            .then(res => res.json())
            .then(data => document.getElementById('response').textContent = data.response || data.error)
            .catch(err => document.getElementById('response').textContent = '錯誤: ' + err);
        }
    </script>
</body>
</html>''')

if __name__ == "__main__":
    create_template()
    if initialize_system():
        app.run(host='0.0.0.0', port=5000, debug=False)
    else:
        logger.error("System initialization failed")