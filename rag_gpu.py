import os
import glob
import PyPDF2
import warnings
from llama_index.core import VectorStoreIndex, Settings, PromptTemplate, Document
from llama_index.core.storage import StorageContext
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama
import chromadb
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging
from flask import Flask, render_template, request, jsonify
import psutil
from docx import Document as DocxDocument
from openpyxl import load_workbook
import re

# ========== 初始化設定 ==========
os.environ["TOKENIZERS_PARALLELISM"] = "false"
warnings.filterwarnings("ignore")

# 精密日誌設定
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.FileHandler('rag_accurate.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# 文件路徑設定
KNOWLEDGE_BASE_PATH = "./knowledge_base"
VECTOR_DB_PATH = "./vector_db"
os.makedirs(KNOWLEDGE_BASE_PATH, exist_ok=True)
os.makedirs(VECTOR_DB_PATH, exist_ok=True)

# ========== 強化文件處理 ==========
def clean_text(text):
    """精密文本清洗"""
    text = re.sub(r'\s+', ' ', text)  # 合併多餘空白
    text = re.sub(r'[\x00-\x1F\x7F]', '', text)  # 移除控制字符
    return text.strip()

def load_documents():
    """高精度文件讀取器"""
    documents = []
    for file_path in glob.glob(os.path.join(KNOWLEDGE_BASE_PATH, '*')):
        try:
            text = ""
            meta = {
                "source": os.path.basename(file_path),
                "type": os.path.splitext(file_path)[1][1:].upper(),
                "pages": ""
            }

            if file_path.lower().endswith('.pdf'):
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    meta["pages"] = str(len(reader.pages))
                    text = "\n\n".join(
                        clean_text(page.extract_text() or f"PDF頁面 {i+1} 無文本內容")
                        for i, page in enumerate(reader.pages[:100])  # 限制100頁
                    )

            elif file_path.lower().endswith('.docx'):
                doc = DocxDocument(file_path)
                text = "\n\n".join(
                    clean_text(para.text) 
                    for para in doc.paragraphs 
                    if clean_text(para.text)
                )

            elif file_path.lower().endswith('.xlsx'):
                wb = load_workbook(file_path)
                for sheet in wb:
                    text += f"\n\n【工作表】{sheet.title}\n"
                    for row in sheet.iter_rows(values_only=True):
                        text += "|".join(
                            str(cell)[:100] if cell else "" 
                            for cell in row
                        ) + "\n"

            elif file_path.lower().endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    text = clean_text(f.read())

            if text:
                documents.append(Document(
                    text=text,
                    metadata=meta
                ))
                logger.info(f"已載入: {meta['source']} (類型: {meta['type']}, 長度: {len(text)}字元)")

        except Exception as e:
            logger.error(f"文件處理失敗: {file_path}\n{str(e)}")

    return documents

# ========== 精準文本分割 ==========
def split_documents(documents):
    """強化中文文本分割"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=512,
        chunk_overlap=128,
        separators=["\n\n", "。", "！", "？", "；", "\n", "，", " "],
        length_function=len,
        is_separator_regex=False
    )

    nodes = []
    for doc in documents:
        chunks = text_splitter.split_text(doc.text)
        for i, chunk in enumerate(chunks):
            chunk_meta = doc.metadata.copy()
            chunk_meta.update({
                "chunk_id": i,
                "char_count": len(chunk),
                "content_hash": hash(chunk)
            })
            nodes.append(Document(
                text=chunk,
                metadata=chunk_meta
            ))
    
    logger.info(f"分割完成: 共 {len(nodes)} 個文本塊")
    return nodes

# ========== 精確回答設定 ==========
def get_precise_prompt_template():
    """高精度提示模板"""
    return PromptTemplate("""
    你是一個專業的繁體中文AI助手，請嚴格按照以下規則回答：

    【上下文】
    {context_str}

    【問題】
    {query_str}

    【回答要求】
    1. 必須使用臺灣正體中文
    2. 答案需直接引用上下文內容
    3. 若上下文無相關資訊，必須回答「根據現有資料無法回答此問題」
    4. 禁止臆測或創造未提及的資訊
    5. 重要數據需標註來源文件

    請開始回答：
    """)

# ========== 系統初始化 ==========
def initialize_system():
    """精密系統初始化"""
    # 嵌入模型設定
    Settings.embed_model = HuggingFaceEmbedding(
        model_name="BAAI/bge-large-zh-v1.5",  # 使用大模型提高精度
        device="cpu",
        embed_batch_size=4  # 小批次確保穩定
    )

    # LLM 設定
    Settings.llm = Ollama(
        model="llama3:8b",
        temperature=0.3,  # 降低隨機性
        request_timeout=600,
        additional_kwargs={
            "options": {
                "num_ctx": 4096,
                "num_thread": 8
            }
        }
    )

    # 向量數據庫
    chroma_client = chromadb.PersistentClient(path=VECTOR_DB_PATH)
    try:
        chroma_client.delete_collection("precise_kb")
    except:
        pass

    vector_store = ChromaVectorStore(
        chroma_collection=chroma_client.create_collection(
            name="precise_kb",
            metadata={
                "hnsw:space": "cosine",
                "precision": "high"
            }
        )
    )
    return StorageContext.from_defaults(vector_store=vector_store)

# ========== Flask 應用 ==========
app = Flask(__name__)
storage_context = initialize_system()

@app.route('/')
def home():
    return render_template('precise_ui.html',
                         ram_available=f"{psutil.virtual_memory().available/1024**3:.1f}GB")

@app.route('/query', methods=['POST'])
def precise_query():
    data = request.get_json()
    if not data or 'query' not in data:
        return jsonify({"error": "無效請求"}), 400

    try:
        # 精確檢索 (增加檢索數量)
        query_engine = index.as_query_engine(
            similarity_top_k=5,
            text_qa_template=get_precise_prompt_template(),
            node_postprocessors=[]
        )
        response = query_engine.query(data['query'].strip())
        
        # 驗證回答質量
        if "無法回答" in str(response) and len(response.source_nodes) > 0:
            response = "找到相關資料但內容不完整，建議查閱原始文件：" + \
                      ", ".join(n.metadata['source'] for n in response.source_nodes)

        return jsonify({
            "response": str(response),
            "sources": [n.metadata['source'] for n in response.source_nodes]
        })

    except Exception as e:
        logger.error(f"查詢失敗: {str(e)}")
        return jsonify({"error": "系統處理異常"}), 500

# ========== 主程序 ==========
if __name__ == "__main__":
    logger.info("===== 高精度RAG系統啟動 =====")
    
    # 1. 精密文件載入
    documents = load_documents()
    if not documents:
        logger.error("錯誤: 未找到可處理文件")
        exit(1)

    # 2. 精準文本處理
    nodes = split_documents(documents)

    # 3. 高精度索引構建
    logger.info("正在建立高精度索引...")
    index = VectorStoreIndex(
        nodes,
        storage_context=storage_context,
        show_progress=True
    )

    # 4. 啟動服務
    logger.info("✅ 系統準備就緒")
    app.run(host='0.0.0.0', port=5000, threaded=False)
